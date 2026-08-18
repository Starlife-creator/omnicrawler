"""S2：ConvertX document 族 —— 文档 → 文本/Markdown 导出（document_ir 桥接）。

txt/eml 仅标准库；docx 依赖 python-docx，缺失时 importorskip。
"""

from __future__ import annotations

from pathlib import Path

import pytest

from omnicrawler.convertx import WRITERS, convert, sniff_format
from omnicrawler.convertx.paths import enumerate_paths


def _make_txt(path: Path) -> Path:
    path.write_text("标题文档\n\n第一段内容。\n\n第二段内容。", encoding="utf-8")
    return path


def _make_eml(path: Path) -> Path:
    content = (
        "From: a@b.com\r\n"
        "Subject: mail subject\r\n"
        "Content-Type: text/plain; charset=utf-8\r\n"
        "\r\n"
        "邮件正文第一段。\r\n\r\n"
        "邮件正文第二段。"
    )
    path.write_bytes(content.encode("utf-8"))
    return path


# ── sniff / 注册 ─────────────────────────────────────────
def test_sniff_document_formats() -> None:
    assert sniff_format(Path("a.txt")) == ".txt"
    assert sniff_format(Path("a.md")) == ".md"
    assert sniff_format(Path("a.docx")) == ".docx"
    assert sniff_format(Path("a.epub")) == ".epub"


def test_document_writers_registered() -> None:
    for ext in (".txt", ".md"):
        assert ext in WRITERS


def test_enumerate_paths_document_family() -> None:
    paths = {f"{p.src_family}->{p.dst_family}": p for p in enumerate_paths()}
    assert "document->document" in paths
    assert paths["document->document"].enabled is True


# ── txt → txt / md（仅标准库）────────────────────────────
def test_convert_txt_to_txt(tmp_path: Path) -> None:
    src = _make_txt(tmp_path / "in.txt")
    dst = tmp_path / "out.txt"
    result = convert(src, dst)
    assert result.target_format == ".txt"
    out = dst.read_text(encoding="utf-8")
    assert "标题文档" in out
    assert "第一段内容。" in out


def test_convert_txt_to_markdown(tmp_path: Path) -> None:
    src = _make_txt(tmp_path / "in.txt")
    dst = tmp_path / "out.md"
    result = convert(src, dst)
    assert result.target_format == ".md"
    md = dst.read_text(encoding="utf-8")
    assert md.startswith("# 标题文档")
    assert "第一段内容。" in md


def test_convert_eml_to_txt(tmp_path: Path) -> None:
    src = _make_eml(tmp_path / "mail.eml")
    dst = tmp_path / "out.txt"
    result = convert(src, dst)
    assert result.rows == 1
    out = dst.read_text(encoding="utf-8")
    assert "mail subject" in out
    assert "邮件正文第一段。" in out


# ── docx → md（依赖 python-docx）─────────────────────────
def test_convert_docx_to_markdown(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    import docx

    src = tmp_path / "report.docx"
    document = docx.Document()
    document.add_heading("经营分析", level=1)
    document.add_paragraph("本季度收入增长。")
    document.save(str(src))

    dst = tmp_path / "out.md"
    result = convert(src, dst)
    assert result.target_format == ".md"
    md = dst.read_text(encoding="utf-8")
    assert md.startswith("# 经营分析")
    assert "本季度收入增长。" in md
